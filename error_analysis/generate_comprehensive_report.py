import os
import sys
import subprocess

def install(pkg):
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q"])

try:
    import docx
except ImportError:
    install("python-docx")
    import docx

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DUAL = os.path.join(SCRIPT_DIR, "error_analysis", "results", "dual_error_histogram.png")
OUT_DOC = os.path.join(SCRIPT_DIR, "Final_Dual_Head_Error_Analysis_Report_v6.docx")

def create_doc():
    doc = docx.Document()
    
    # Title
    title = doc.add_heading('Dual-Head Error Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ── 1. Methodology ──────────────────────────────────────────────────────────
    doc.add_heading('1. Contrastive Error Analysis Methodology', level=1)
    doc.add_paragraph(
        "To rigorously evaluate the DeBERTa model's multi-head performance beyond standard aggregate metrics, "
        "we implemented a Contrastive Error Analysis pipeline. Rather than looking at aggregate failure rates, "
        "this method systematically matches incorrect predictions with successful predictions that share the exact same metadata "
        "(e.g., ground truth label, scenario, and communication style). These matched 'Incorrect-Correct Pairs' (ICPs) "
        "allow us to isolate specific linguistic, structural, or contextual features, enabling us to identify the underlying "
        "algorithmic blindspots confusing the model."
    )

    # ── 2. Accuracy Comparison ─────────────────────────────────────────────────
    doc.add_heading('2. Dual-Head Accuracy Comparison', level=1)
    doc.add_paragraph(
        "Evaluating the model against the 750 samples in the test set reveals a clear performance discrepancy "
        "between the two classification heads:"
    )
    
    p = doc.add_paragraph()
    p.add_run("Urgency Head Errors: ").bold = True
    p.add_run("266 mistakes (~64.5% accuracy)\n")
    p.add_run("Emotion Head Errors: ").bold = True
    p.add_run("171 mistakes (~77.2% accuracy)")
    
    doc.add_paragraph(
        "The DeBERTa model performs significantly better at classifying Emotion than Urgency. "
        "Urgency is inherently more difficult to classify as it requires deep contextual understanding of "
        "socio-economic impact, deadlines, and technical severity. Emotion, while nuanced, often has "
        "more overt vocabulary triggers."
    )
    
    if os.path.exists(IMG_DUAL):
        doc.add_picture(IMG_DUAL, width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figure 1: Dual-Head Error Distribution (Urgency vs Emotion).").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Dual Histogram Image Missing]")

    # ── 3. Urgency Head Analysis ───────────────────────────────────────────────
    doc.add_heading('3. Urgency Head Analysis', level=1)
    doc.add_paragraph(
        "Out of 266 total urgency mistakes, 140 were successfully paired into ICPs for contrastive qualitative analysis."
    )


    doc.add_heading('Top 3 Urgency Blindspots', level=2)
    table_u = doc.add_table(rows=1, cols=2)
    table_u.style = 'Table Grid'
    hdr_cells_u = table_u.rows[0].cells
    hdr_cells_u[0].text = 'Algorithmic Blindspot'
    hdr_cells_u[1].text = 'Description'
    
    urgency_blindspots = [
        ("Emotional & Vulnerability Over-indexing", "The model frequently misinterprets expressions of emotional distress or personal vulnerability as 'High' urgency, while failing to adequately weigh explicit statements of patience or willingness to wait."),
        ("Contextual Misinterpretation of Deadlines", "The model over-prioritizes explicit deadlines or external escalation threats (e.g., Ombudsman, legal action) without sufficiently contextualizing the actual impact or immediacy of the threat."),
        ("Failure to Synthesize Problem History", "The model struggles to weigh the cumulative effect of repeated service failures or prolonged unresolved issues, over-indexing on single recent events instead of persistent crises.")
    ]
    
    for bs, desc in urgency_blindspots:
        row_cells = table_u.add_row().cells
        row_cells[0].text = bs
        row_cells[1].text = desc

    doc.add_paragraph("\n")

    # ── 4. Emotion Head Analysis ───────────────────────────────────────────────
    doc.add_heading('4. Emotion Head Analysis', level=1)
    doc.add_paragraph(
        "Out of 171 total emotion mistakes, 111 were successfully paired into ICPs for contrastive qualitative analysis."
    )

    doc.add_heading('Top 3 Emotion Blindspots', level=2)
    table_e = doc.add_table(rows=1, cols=2)
    table_e.style = 'Table Grid'
    hdr_cells_e = table_e.rows[0].cells
    hdr_cells_e[0].text = 'Algorithmic Blindspot'
    hdr_cells_e[1].text = 'Description'
    
    emotion_blindspots = [
        ("Misinterpreting Cumulative Impact", "The model struggles to infer emotion (Low, Medium, High) from the cumulative weight of repeated issues or severe financial strain without explicit high-impact angry keywords."),
        ("Over-reliance on Overt Aggressive Language", "The model over-indexes on explicit personal frustration or aggressive 'power phrases', misclassifying intensity when the tone is strictly formal but legally threatening."),
        ("Under-recognition of Internal Escalation", "The model fails to identify 'Medium' or 'High' emotional severity in initial contacts due to critical system failures without explicit threats.")
    ]
    
    for bs, desc in emotion_blindspots:
        row_cells = table_e.add_row().cells
        row_cells[0].text = bs
        row_cells[1].text = desc

    doc.add_paragraph("\n")

    # ── 5. Adversarial Testing ─────────────────────────────────────────────────
    doc.add_heading('5. Adversarial Stress Testing', level=1)
    doc.add_paragraph(
        "To actively validate the blindspots diagnosed via the contrastive analysis, we conducted a 10-item Adversarial Test. "
        "Unlike the passive contrastive analysis which relies on naturally occurring test set data, the adversarial test "
        "uses synthetic, edge-case perturbations designed specifically to trick the model. The model passed 6/10 tests, but "
        "crucially, the specific failures mapped perfectly to our established contrastive blindspots."
    )
    
    doc.add_paragraph(
        "For example, the adversarial test 'Passive-Aggressive Legal Threat' (True Emotion: High, Predicted: Low) failed exactly as predicted "
        "by Emotion Blindspot #2 (Over-reliance on Overt Aggressive Language). It lacked explicit angry words, but carried extreme threat. "
        "Similarly, the 'Sarcastic Praise' test failed because the model misinterpreted the cumulative impact (Emotion Blindspot #1) "
        "when masked by mathematically positive words. These complementary results confirm the extreme robustness of our passive error analysis."
    )

    # ── 6. Limitations ─────────────────────────────────────────────────────────
    doc.add_heading('6. Limitations of the Approach', level=1)
    doc.add_paragraph(
        "Based on the methodology used to create the original dataset, several inherent limitations "
        "must be acknowledged regarding the model's training and evaluation:"
    )
    
    LIMITATIONS = [
        ("Synthetic Data Bias", "The entire dataset of 5,000 complaints was synthetically generated using OpenAI's GPT-4o-mini. Consequently, the DeBERTa model is primarily learning the linguistic patterns, vocabulary, and rhetorical structures typical of Large Language Models (LLMs) rather than genuine human customers. Real-world complaints often display far more chaotic and unpredictable sentence structures."),
        ("Strict Affinity Constraints", "The dataset generation enforced strict logical rules (e.g., a 'Complete Service Outage' could never be assigned 'Low Urgency', and 'Passive-aggressive/sarcastic' styles could never be assigned 'Low Emotion'). While this ensures clean synthetic data, in the real world, customers can write highly contradictory or idiosyncratic complaints that violate these logical boundaries."),
        ("Constrained Diversity", "Although varied, the dataset is still ultimately constrained to 20 scenarios, 8 writing styles, and 8 customer profiles. True real-world telecommunication complaints span an infinitely broader spectrum of unique personal situations and sub-scenarios."),
        ("Lack of Authentic Noise", "The synthetic text inherently lacks the authentic, natural noise found in live production environments—such as severe spelling mistakes, SMS abbreviations, and stream-of-consciousness formatting. Our adversarial testing partially mitigates this, but cannot fully replicate true production noise.")
    ]
    
    for title, desc in LIMITATIONS:
        p = doc.add_paragraph()
        p.add_run(title + ": ").bold = True
        p.add_run(desc)
        
    # ── 7. Further Improvements ────────────────────────────────────────────────
    doc.add_heading('7. Further Improvements', level=1)
    doc.add_paragraph(
        "To enhance the robustness of this dual-head architectural approach, future iterations must systematically "
        "transition from synthetic validation to real-world deployment evaluation while conducting targeted data augmentation. "
        "Although the DeBERTa model demonstrates strong baseline capabilities, its exclusive reliance on GPT-4o-mini generated "
        "data restricts its exposure to genuine linguistic noise, spelling errors, and idiosyncratic formatting found in live "
        "production environments. Furthermore, our contrastive error analysis exposed rigid algorithmic blindspots, such as "
        "the model's over-reliance on overt aggressive language and its consistent misinterpretation of cumulative problem history. "
        "Therefore, validating the model against a 'Gold Standard' real-world telecommunication complaint dataset is essential "
        "to measure its true generalizability. Concurrently, implementing a data augmentation strategy that specifically injects "
        "adversarial edge-cases based on our identified errors—such as formal legal threats conveying high emotion—will directly "
        "re-calibrate the model's distorted decision boundaries, ensuring optimal reliability in a live triage pipeline."
    )

    doc.save(OUT_DOC)
    print(f"Document saved successfully to {OUT_DOC}")

if __name__ == "__main__":
    create_doc()
